#!/usr/bin/env python3
"""
LoadGuard Onboarding Document Generator
Generates DOCX offer letters, equity incentive letters, vesting schedules,
task schedules, and immigration support letters for new joiners.

Usage:
  python3 .claude/scripts/generate_onboard_docs.py --variant A1 --name "John Doe" ...

Variants:
  A1 - US Citizen, Equitied Employee
  A2 - US Citizen, Task-Compensated Employee
  B1 - Non-US Citizen, Equitied Employee (Immigration-Ready)
  B2 - Non-US Citizen, Task-Compensated Employee (Immigration-Ready)
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

import argparse
import csv
import json
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Constants ───────────────────────────────────────────────────────────────

COMPANY_NAME = "Project LoadGuard Incorporated"
COMPANY_ADDR = "123 Example Street"
COMPANY_CITY = "Seattle, WA 98101"
COMPANY_EIN = "XX-XXXXXXX"
CEO_NAME = "Jane Doe"
CEO_TITLE = "Chief Executive Officer & President"
ACCENT_COLOR = "C02B2B"
ACCENT_RGB = RGBColor(0xC0, 0x2B, 0x2B)
GRAY_RGB = RGBColor(0x55, 0x55, 0x55)
BODY_RGB = RGBColor(0x33, 0x33, 0x33)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, '..', '..'))
LOGO_PATH = os.path.join(PROJECT_ROOT, 'lg-logo.png.png')

DEFAULT_MILESTONES = [
    {"description": "Prototype working & deployable (v1 functional build)", "shares": None},
    {"description": "Device detects anomalies with 60%+ accuracy, 3-nines reliability", "shares": None},
    {"description": "Full voyage week test pass, power/boot/comms proven (v2)", "shares": None},
    {"description": "Live pilot with paying partner, proven alerting & service uptime", "shares": None},
    {"description": "Device mass-manufacturable with optimized design", "shares": None},
    {"description": "Investment round sourced from VC or corporate VC", "shares": None},
    {"description": "Manufacturing begins at scale with 6-sigma standard", "shares": None},
    {"description": "Scales to 10+ paying customers with proven robustness", "shares": None},
]


# ─── Helpers ─────────────────────────────────────────────────────────────────

def create_doc_with_letterhead():
    """Create a new Document with LoadGuard letterhead."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = BODY_RGB
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    # Header
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()

    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = True

    tbl = header_table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        borders = tblPr.find(qn('w:tblBorders'))
        if borders is not None:
            tblPr.remove(borders)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Logo
    logo_cell = header_table.cell(0, 0)
    logo_cell.width = Inches(1.5)
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_PATH):
        run = logo_para.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.0))

    # Company info
    info_cell = header_table.cell(0, 1)
    info_cell.width = Inches(5.0)
    info_para = info_cell.paragraphs[0]
    info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    lines = [
        (COMPANY_NAME, True, Pt(12)),
        (COMPANY_ADDR, False, Pt(9)),
        (COMPANY_CITY, False, Pt(9)),
        (f"EIN: {COMPANY_EIN}", False, Pt(9)),
    ]
    for i, (text, bold, size) in enumerate(lines):
        run = info_para.add_run(text)
        run.bold = bold
        run.font.size = size
        run.font.name = 'Calibri'
        run.font.color.rgb = ACCENT_RGB if bold else GRAY_RGB
        if i < len(lines) - 1:
            info_para.add_run("\n").font.size = Pt(4)

    # Red accent line
    border_para = header.add_paragraph()
    border_para.paragraph_format.space_before = Pt(6)
    border_para.paragraph_format.space_after = Pt(0)
    pPr = border_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), ACCENT_COLOR)
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Remove initial empty paragraph
    if header.paragraphs[0].text == '' and len(header.paragraphs[0].runs) == 0:
        p_element = header.paragraphs[0]._element
        p_element.getparent().remove(p_element)

    return doc


def add_body(doc, text, space_after=Pt(10), bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.bold = bold
    return p


def add_section(doc, number, title, body_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)

    run = p.add_run(f"{number}.")
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.bold = True

    run = p.add_run(f"\t{title}.  ")
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.bold = True

    run = p.add_run(body_text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p


def add_signature_block(doc, include_employee_name=""):
    doc.add_paragraph()
    sincerely = doc.add_paragraph()
    sincerely.paragraph_format.space_after = Pt(36)
    run = sincerely.add_run("Sincerely,")
    run.font.size = Pt(11)

    for text, bold in [
        (CEO_NAME, True),
        (CEO_TITLE, False),
        (COMPANY_NAME, False),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)

    for label in ["Signature: ______________________________", "Date: ______________________________"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2) if "Signature" in label else Pt(30)
        run = p.add_run(label)
        run.font.size = Pt(11)

    if include_employee_name:
        accept = doc.add_paragraph()
        accept.paragraph_format.space_before = Pt(12)
        accept.paragraph_format.space_after = Pt(12)
        run = accept.add_run("ACCEPTED AND AGREED:")
        run.bold = True
        run.font.size = Pt(11)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(include_employee_name)
        run.bold = True
        run.font.size = Pt(11)

        for label in ["Signature: ______________________________", "Date: ______________________________"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2) if "Signature" in label else Pt(6)
            run = p.add_run(label)
            run.font.size = Pt(11)


def today_str():
    return datetime.now().strftime("%B %d, %Y").replace(" 0", " ")


def ensure_output_dir(name):
    parts = name.strip().split()
    first = parts[0]
    last = parts[-1] if len(parts) > 1 else parts[0]
    folder = os.path.join(PROJECT_ROOT, "Admin", "People", f"{first} {last}")
    os.makedirs(folder, exist_ok=True)
    return folder


# ─── Document Generators ────────────────────────────────────────────────────

def generate_offer_letter(args, immigration=False, task_based=False):
    """Generate the offer of employment letter (lawyer template language)."""
    doc = create_doc_with_letterhead()
    is_immigration = immigration
    name = args.name
    position = args.position
    first_name = name.split()[0]
    is_part_time = getattr(args, 'part_time', False)
    emp_label = "part-time" if is_part_time else "full-time"
    has_tasks = task_based and args.tasks
    has_salary = args.salary and args.salary.upper() != "TBD"

    # Date
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(today_str())
    run.font.size = Pt(11)

    # Addressee
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(name)
    run.font.size = Pt(11)

    # Re line
    re_text = f"Re: Offer of {emp_label.title()} Employment \u2014 {position}"
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(re_text)
    run.bold = True
    run.font.size = Pt(11)

    # Salutation
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"Dear {first_name}:")
    run.font.size = Pt(11)

    # Intro
    intro = (
        f"{COMPANY_NAME} (the \u201cCompany\u201d) is pleased to offer you "
        f"employment with the Company on the terms described below."
    )
    if is_immigration:
        intro += (
            " This letter shall serve as an official offer of "
            f"{emp_label} employment and may be presented to U.S. Citizenship "
            "and Immigration Services (\u201cUSCIS\u201d), the U.S. Department of "
            "Labor (\u201cDOL\u201d), or any other government agency as evidence of "
            f"a bona fide offer of {emp_label} employment in a specialty occupation."
        )
    add_body(doc, intro, space_after=Pt(14))

    # ── Sections ──
    sec = 1

    # 1. Position (+ work location folded in)
    pos_text = (
        f"You will start in a {emp_label} position as {position} and you will "
        f"initially report to the Company\u2019s {args.reports_to}."
    )
    if is_part_time:
        hours_spec = getattr(args, 'hours_per_week', None)
        if hours_spec:
            pos_text += (
                f"  As a part-time employee, you will work approximately "
                f"{hours_spec} hours a week, although your actual hours worked "
                "may change from time to time based on the then current needs "
                "of the Company."
            )
        else:
            pos_text += (
                "  As a part-time employee, your hours will be determined "
                "based on the then current needs of the Company."
            )
    if is_immigration:
        pos_text += (
            "  This is a specialty occupation position requiring the theoretical "
            "and practical application of a body of highly specialized knowledge. "
            "The minimum educational requirement for this position is a "
            "bachelor\u2019s degree (or its foreign equivalent) in Computer Science, "
            "Computer Engineering, Electrical Engineering, or a directly related "
            "field."
        )
    # Work location
    work_loc = args.work_location
    if work_loc and work_loc.lower() == "remote":
        pos_text += (
            "  Your primary work location will be remote, or at such other "
            "location as the Company may designate from time to time"
        )
    elif work_loc:
        pos_text += (
            f"  Your primary work location will be at the Company\u2019s offices "
            f"at {work_loc}, or at such other location as the Company may "
            "designate from time to time"
        )
    if is_immigration and work_loc:
        pos_text += (
            " in accordance with applicable immigration regulations.  Any "
            "change in work location will comply with applicable H-1B "
            "requirements, including the filing of an amended petition or "
            "new Labor Condition Application if required."
        )
    elif work_loc:
        pos_text += "."
    add_section(doc, sec, "Position" + (" (Specialty Occupation)" if is_immigration else ""), pos_text)
    sec += 1

    # 2. Introductory Period (optional)
    if getattr(args, 'introductory_period', False):
        intro_days = getattr(args, 'intro_days', 30)
        ip_text = (
            f"Your first {intro_days} calendar days of employment shall "
            f"constitute an Introductory Period. During the Introductory "
            f"Period, you will be a {emp_label} employee of the Company with "
            "all the rights, compensation, and benefits set forth in this "
            "letter and required by applicable law. The purpose of the "
            "Introductory Period is to provide structured onboarding and to "
            "allow both you and the Company to assess mutual fit and alignment "
            "of expectations. Upon successful completion of the Introductory "
            "Period, your employment will continue under the same terms set "
            "forth herein without interruption. For the avoidance of doubt, "
            "the Introductory Period is a component of the "
            f"{emp_label} employment relationship described in this letter and "
            "does not constitute a separate, preliminary, or contingent "
            "engagement."
        )
        if is_immigration:
            ip_text += (
                " Your compensation during the Introductory Period shall not "
                "be less than the applicable prevailing wage."
            )
        ip_text += (
            " The Introductory Period does not alter the at-will nature of "
            "your employment; either party may terminate the employment "
            "relationship at any time during or after the Introductory Period, "
            "with or without cause or notice."
        )
        add_section(doc, sec, "Introductory Period", ip_text)
        sec += 1

    # 3. Base Compensation
    if has_tasks:
        if is_immigration:
            # B2 with tasks: base salary + task bonuses
            salary_str = args.salary if has_salary else "[BASE SALARY]"
            comp_text = (
                f"You will be paid a starting salary at the rate of "
                f"{salary_str}, payable in accordance with the Company\u2019s "
                "standard payroll schedule.  In addition to your base salary, "
                "you will be eligible for task completion bonuses as set forth "
                "in the Task Bonus Schedule attached hereto as Exhibit A.  "
                "The Company confirms that your base salary alone will be at "
                "or above the prevailing wage for the "
                f"{position} position at the designated work location, as "
                "determined by the applicable wage source and as required by "
                "the U.S. Department of Labor and all regulations governing "
                "H-1B nonimmigrant workers."
            )
        else:
            # A2: task-based compensation
            comp_text = (
                "Your compensation will be determined on a task-completion "
                "basis as set forth in the Task Schedule attached hereto as "
                "Exhibit A. Each task defines a scope of work and a "
                "corresponding payment amount. Payment will be made upon the "
                "Company\u2019s confirmation that the deliverable has been "
                "satisfactorily completed, in accordance with the Company\u2019s "
                "standard payroll schedule."
            )
    else:
        # Salary-based (A1, B1, or B2 with salary and no tasks)
        if has_salary:
            comp_text = (
                f"You will be paid a starting salary at the rate of "
                f"{args.salary}, payable in accordance with the Company\u2019s "
                "standard payroll schedule."
            )
        else:
            comp_text = (
                "Your compensation, including base salary, benefits, and any "
                "other terms of remuneration, will be set forth in a separate "
                "compensation agreement between you and the Company."
            )
        if is_immigration:
            comp_text += (
                f"  The Company confirms that your compensation will be at or "
                f"above the prevailing wage for the {position} position at the "
                "designated work location, as determined by the applicable "
                "wage source and as required by the U.S. Department of Labor "
                "and all regulations governing H-1B nonimmigrant workers."
            )
        variant = getattr(args, 'variant', '').upper()
        if variant in ('A1', 'B1'):
            comp_text += (
                "  Additional details regarding equity compensation, if "
                "applicable, are set forth in a separate Equity Incentive "
                "Letter."
            )
    add_section(doc, sec, "Base Compensation", comp_text)
    sec += 1

    # 4. Taxes, Withholding and Required Deductions (NEW)
    add_section(doc, sec, "Taxes, Withholding and Required Deductions",
        "All forms of compensation referred to in this letter are subject to "
        "all applicable taxes, withholding and any other deductions required "
        "by applicable law."
    )
    sec += 1

    # 5. Employee Benefits (lawyer template + part-time language)
    ben_text = (
        "As a regular employee of the Company, you will be eligible to "
        "participate in the employee benefit plans and programs currently "
        "and hereafter maintained by the Company and generally available to "
        "similarly situated employees of the Company, subject in each case "
        "to the terms and conditions of the plan in question, including any "
        "eligibility requirements set forth therein, and the determination "
        "of any person or committee administering the plan.  Notwithstanding "
        "the foregoing, the Company reserves the right to modify or terminate "
        "benefits from time to time as it deems necessary or appropriate in "
        "its sole discretion."
    )
    if is_part_time:
        ben_text += (
            "  As a part-time employee, you may be ineligible for some of "
            "the benefits available to the Company\u2019s full-time employees.  "
            "Your eligibility will be governed by the Company\u2019s applicable "
            "policies and plans.  You will be afforded the benefits mandated "
            "by law, including paid sick leave (if applicable) and workers\u2019 "
            "compensation insurance."
        )
    if getattr(args, 'introductory_period', False):
        ben_text += (
            " Eligibility for benefits begins on your first day of "
            "employment, including during the Introductory Period."
        )
    add_section(doc, sec, "Employee Benefits", ben_text)
    sec += 1

    # 6. Start Date
    sd_text = "Your anticipated start date "
    if args.start_date and args.start_date.upper() != "TBD":
        sd_text += f"is {args.start_date}"
    else:
        sd_text += "will be communicated to you in writing"
    if is_immigration:
        sd_text += (
            " and is contingent upon your obtaining the necessary work "
            "authorization to be lawfully employed in the United States, "
            "including but not limited to approval of an H-1B petition (or "
            "change of status or other applicable visa petition) filed on "
            "your behalf by the Company."
        )
    else:
        sd_text += "."
    add_section(doc, sec, "Start Date", sd_text)
    sec += 1

    # 7. Immigration Sponsorship (B variants only)
    if is_immigration:
        imm_text = (
            "The Company intends to file an H-1B petition (or other "
            "applicable visa petition) with USCIS on your behalf for the "
            f"{position} position described herein. The Company will bear "
            "the costs associated with the filing of the H-1B petition, "
            "including attorney fees and filing fees, as required by law. "
            "The Company will file the required Labor Condition Application "
            "(\u201cLCA\u201d) with the DOL and will comply with all LCA obligations, "
            "including payment of the prevailing wage or the actual wage, "
            "whichever is higher. This offer of employment is contingent "
            "upon your obtaining and maintaining valid work authorization in "
            "the United States. The Company will cooperate in good faith "
            "with all immigration-related processes and will comply with all "
            "applicable immigration laws and regulations."
        )
        add_section(doc, sec, "Immigration Sponsorship", imm_text)
        sec += 1

    # 8. At-Will Employment (lawyer template exact language)
    aw_text = (
        "Employment with the Company is for no specific period of time.  "
        "Your employment with the Company will be \u201cat will,\u201d meaning that "
        "either you or the Company may terminate your employment at any time "
        "and for any reason, with or without cause or notice.  Any contrary "
        "representations which may have been made to you are superseded by "
        "this offer.  This is the full and complete agreement between you "
        "and the Company on this term.  Although your job duties, title, "
        "compensation and benefits, as well as the Company\u2019s personnel "
        "policies and procedures, may change from time to time, the \u201cat "
        "will\u201d nature of your employment may only be changed in an express "
        "written agreement signed by you and the Company\u2019s Chief Executive "
        "Officer."
    )
    if is_immigration:
        aw_text += (
            "  In the event the Company terminates your employment, the "
            "Company will comply with all obligations under applicable "
            "immigration law, including but not limited to offering to pay "
            "reasonable costs of return transportation to your last country "
            "of residence, as required under 8 C.F.R. \u00a7 214.2(h)(4)(iii)(E)."
        )
    add_section(doc, sec, "At-Will Employment", aw_text)
    sec += 1

    # 9. Equity Forfeiture and Repurchase (A1/B1 only)
    variant = getattr(args, 'variant', '').upper()
    if variant in ('A1', 'B1'):
        forfeit_text = (
            "You acknowledge and agree that any equity grant described in a "
            "separate Equity Incentive Letter or stock purchase agreement is "
            "subject to the following terms, which shall survive termination "
            "of your employment:"
            "\n\n"
            "(a) Forfeiture of Unvested Shares. Upon termination of your "
            "employment for any reason (whether voluntary or involuntary, "
            "with or without Cause), all unvested shares\u2014including shares "
            "subject to unmet time-based vesting conditions and shares tied "
            "to milestones not yet achieved and confirmed by the Board\u2014shall "
            "be immediately and automatically forfeited to the Company for "
            "no consideration."
            "\n\n"
            "(b) Company Repurchase Right on Vested Shares. Upon termination "
            "of your employment, the Company shall have the right (but not "
            "the obligation) to repurchase any or all of your vested shares "
            "at a price equal to (i) the fair market value as determined by "
            "the Board in good faith if termination is without Cause, or "
            "(ii) the lower of the original purchase price per share or fair "
            "market value if termination is for Cause. The Company may "
            "exercise this right within 180 days of termination by delivering "
            "written notice to you."
            "\n\n"
            "(c) Clawback. If the Company determines that any equity was "
            "granted or vested based on fraud, material misrepresentation, or "
            "breach of the CIIAA, the Company may recover such equity (or "
            "the proceeds thereof) from you."
            "\n\n"
            "(d) No Guaranteed Right to Equity. You acknowledge that any "
            "equity described herein or in a separate Equity Incentive Letter "
            "reflects management\u2019s recommendation to the Board and does not "
            "constitute a binding promise or guarantee of any equity grant. "
            "The Board retains sole discretion to approve, modify, or decline "
            "any equity grant."
        )
        add_section(doc, sec, "Equity Forfeiture and Repurchase Rights", forfeit_text)
        sec += 1

    # 10. CIIAA (Attachment A — lawyer template language)
    add_section(doc, sec,
        "Confidential Information and Invention Assignment Agreement",
        "You are required, as a condition of your employment with the "
        "Company, to sign the Company\u2019s standard Confidential Information "
        "and Invention Assignment Agreement, a copy of which is attached "
        "hereto as Attachment\u00a0A (the \u201cCIIAA\u201d)."
    )
    sec += 1

    # 11. Company Policies (lawyer template language)
    add_section(doc, sec, "Company Policies",
        "As a condition of your employment, you must become familiar with "
        "all policies, practices and procedures of the Company that are "
        "applicable to you, as such policies are changed from time to time, "
        "and must comply with all such policies, and in each case you agree "
        "that you will do so."
    )
    sec += 1

    # 12. No Conflicts (lawyer template comprehensive version)
    add_section(doc, sec, "No Conflicts",
        "You represent and warrant to the Company that you are under no "
        "obligations or commitments, whether contractual or otherwise, that "
        "would prohibit or otherwise restrict you from performing your "
        "duties for the Company.  Without limiting any terms of the CIIAA, "
        "you may not use or disclose in connection with your performance of "
        "your duties for the Company any trade secrets or other proprietary "
        "information or intellectual property in which you or any other "
        "person (including but not limited to any former employer or company "
        "for whom you consulted) has an interest and you confirm that your "
        "employment with the Company will not infringe or otherwise violate "
        "any other person\u2019s rights.  Further, you must abide by any "
        "contractual obligations to which you are subject that require you "
        "to refrain from soliciting any person employed by or otherwise "
        "associated with any former or current employer or company for whom "
        "you consulted.  You represent and warrant to the Company that you "
        "have returned all property and confidential information belonging "
        "to any prior employer."
    )
    sec += 1

    # 13. Outside Activities (NEW — lawyer template language)
    oa_text = ""
    if not is_part_time:
        oa_text = (
            "While employed by the Company, you must devote your full "
            "business efforts and time to the Company.  "
        )
    oa_text += (
        "You understand you have a duty of loyalty to the Company during "
        "the course of your employment. Without limiting the generality of "
        "the foregoing, to the fullest extent permitted under applicable "
        "laws, while you render services to the Company, you may not engage "
        "in, launch, or encourage others to launch any other company, "
        "venture, employment, consulting project or other business activity "
        "(whether on a full- or part-time basis) that would create a "
        "conflict of interest with the Company or that would, directly or "
        "indirectly, constitute your engagement in or participation in any "
        "business that is competitive in any manner with the Company\u2019s "
        "business."
    )
    add_section(doc, sec, "Outside Activities", oa_text)
    sec += 1

    # 14. Arbitration Agreement (NEW — lawyer template language)
    add_section(doc, sec, "Arbitration Agreement",
        "You and the Company agree that to the fullest extent permitted by "
        "law, any and all claims relating to, arising from or regarding "
        "this offer of employment or your employment with the Company, "
        "including claims by the Company, claims against the Company and "
        "claims against any current or former officer, director, agent, or "
        "employee of the Company, shall be resolved by final and binding "
        "arbitration as set forth in the Mutual Arbitration Agreement, "
        "attached hereto as Attachment\u00a0B (the \u201cArbitration Agreement\u201d), "
        "which you will be expected to sign and return when you return an "
        "executed copy of this offer letter."
    )
    sec += 1

    # 15. Governing Law (lawyer template — state-where-employee-resides)
    add_section(doc, sec, "Governing Law",
        "Except for the Arbitration Agreement, the validity, interpretation, "
        "construction and performance of this letter agreement, and all acts "
        "and transactions pursuant hereto and the rights and obligations of "
        "the parties hereto shall be governed, construed and interpreted in "
        "accordance with the laws of the state in which you are employed by "
        "the Company, or, if you are a remote employee, reside, or in the "
        "event the offer of employment is withdrawn, then in accordance "
        "with the laws of the state in which you resided at the time this "
        "offer was made, without giving effect to principles of conflicts "
        "of law."
    )
    sec += 1

    # 16. Entire Agreement (lawyer template — includes Arbitration Agreement)
    ea_text = (
        "This letter, together with the CIIAA and Arbitration Agreement, "
        "sets forth the entire agreement and understanding of the parties "
        "relating to the subject matter herein and supersedes all prior or "
        "contemporaneous discussions, understandings and agreements, whether "
        "oral or written, between them relating to the subject matter hereof."
    )
    if getattr(args, 'introductory_period', False):
        ea_text += (
            "  For the avoidance of doubt, the Introductory Period described "
            "herein is a component of the employment relationship and does "
            "not constitute a separate or preliminary engagement."
        )
    add_section(doc, sec, "Entire Agreement", ea_text)
    sec += 1

    # 17. Severability (lawyer template exact language)
    add_section(doc, sec, "Severability",
        "If any provision of this letter becomes or is deemed invalid, "
        "illegal or unenforceable in any applicable jurisdiction by reason "
        "of the scope, extent or duration of its coverage, then such "
        "provision shall be deemed amended to the minimum extent necessary "
        "to conform to applicable law so as to be valid and enforceable or, "
        "if such provision cannot be so amended without materially altering "
        "the intention of the parties, then such provision shall be stricken "
        "and the remainder of this letter shall continue in full force and "
        "effect.  If any provision of this letter is rendered illegal by "
        "any present or future statute, law, ordinance or regulation then "
        "that provision shall be curtailed or limited only to the minimum "
        "extent necessary to bring the provision into compliance with "
        "applicable laws.  All the other terms and provisions of this letter "
        "shall continue in full force and effect without impairment or "
        "limitation."
    )
    sec += 1

    # 18. No Assignment (NEW — lawyer template language)
    add_section(doc, sec, "No Assignment",
        "This letter and all of your rights and obligations hereunder are "
        "personal to you and may not be transferred or assigned by you at "
        "any time.  The Company may assign its rights under this letter "
        "freely without restriction, including to any entity that assumes "
        "the Company\u2019s obligations hereunder in connection with any sale or "
        "transfer of all or a substantial portion of the Company\u2019s assets "
        "to such entity."
    )
    sec += 1

    # 19. Counterparts (NEW — lawyer template language)
    add_section(doc, sec, "Counterparts",
        "This letter may be executed in any number of counterparts, each "
        "of which when so executed and delivered shall be deemed an "
        "original, and all of which together shall constitute one and the "
        "same agreement.  Execution via an electronic signature platform or "
        "scanned image will have the same force and effect as execution of "
        "an original, and an electronic signature or scanned image signature "
        "will be deemed an original and valid signature."
    )
    sec += 1

    # 20. Electronic Delivery (NEW — lawyer template language)
    add_section(doc, sec, "Electronic Delivery",
        "The Company may, in its sole discretion, decide to deliver any "
        "documents or notices related to this letter, securities of the "
        "Company or any of its affiliates or any other matter, including "
        "documents and/or notices required to be delivered to you by "
        "applicable securities law or any other law or the Company\u2019s "
        "Certificate of Incorporation or Bylaws by email or any other "
        "electronic means.  You hereby consent to (i)\u00a0conduct business "
        "electronically (ii)\u00a0receive such documents and notices by such "
        "electronic delivery and (iii)\u00a0sign documents electronically and "
        "agree to participate through an on-line or electronic system "
        "established and maintained by the Company or a third party "
        "designated by the Company."
    )

    # ── Closing (lawyer template language) ──
    doc.add_paragraph()
    closing_text = (
        "If you wish to accept this offer, please sign and date this "
        "letter and the enclosed CIIAA and Arbitration Agreement and return "
        "them to me.  As required by law, your employment with the Company "
        "is also contingent upon your providing legal proof of your identity "
        "and authorization to work in the United States"
    )
    if is_immigration:
        closing_text += " (Form I-9 verification)"
    closing_text += (
        ".  In addition, the Company reserves the right to conduct "
        "background investigations and/or reference checks on all of its "
        "potential employees.  Your job offer, therefore, may be contingent "
        "upon a clearance of such a background investigation and/or "
        "reference check, if any."
    )
    add_body(doc, closing_text, space_after=Pt(10))

    doc.add_paragraph()
    add_body(doc, (
        "We look forward to your favorable reply and to working with you "
        f"at {COMPANY_NAME}!"
    ), space_after=Pt(10))

    # Anticipated Start Date
    if args.start_date and args.start_date.upper() != "TBD":
        add_body(doc, f"Anticipated Start Date:  {args.start_date}",
                 space_after=Pt(14), bold=True)

    # Enclosures
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Enclosure:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    add_body(doc, "Attachment A:  Confidential Information and Invention "
             "Assignment Agreement", space_after=Pt(2))
    add_body(doc, "Attachment B:  Mutual Arbitration Agreement",
             space_after=Pt(14))

    # Signature block
    add_signature_block(doc, include_employee_name=name)

    # ── Attachment A page ──
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("ATTACHMENT A")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        "CONFIDENTIAL INFORMATION AND\nINVENTION ASSIGNMENT AGREEMENT"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(See Attached)")
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.italic = True

    # ── Attachment B page ──
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("ATTACHMENT B")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("MUTUAL ARBITRATION AGREEMENT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(See Attached)")
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.italic = True

    return doc

def generate_equity_incentive_letter(args):
    """Generate the equity incentive letter (A1/B1)."""
    doc = create_doc_with_letterhead()
    name = args.name
    first_name = name.split()[0]
    total = args.total_shares or "[TOTAL SHARES]"
    time_shares = args.time_shares or "[TIME-BASED SHARES]"
    milestone_shares = args.milestone_shares or "[MILESTONE-BASED SHARES]"
    cliff = args.cliff_months or 12
    vest_years = args.vest_years or 4
    vest_months = vest_years * 12

    # Date
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(today_str())
    run.font.size = Pt(11)

    # Addressee
    add_body(doc, name, space_after=Pt(18))

    # Re line
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Re: Equity Incentive Grant")
    run.bold = True
    run.font.size = Pt(11)

    add_body(doc, f"Dear {first_name}:", space_after=Pt(12))

    # Intro
    add_body(doc, (
        f"This letter (the \u201cEquity Letter\u201d) is intended to memorialize the terms "
        f"of your equity grant with {COMPANY_NAME} (the \u201cCompany\u201d). This Equity "
        "Letter does not alter your employment relationship with the Company and "
        "does not amend any existing employment documents except as expressly stated "
        "herein."
    ), space_after=Pt(14))

    # Grant Summary
    add_body(doc, "Grant Summary", space_after=Pt(6), bold=True)
    add_body(doc, (
        f"Subject to the approval of the Company\u2019s Board of Directors (the \u201cBoard\u201d), "
        f"management recommends that you be granted {total} shares of Common Stock "
        "of the Company (the \u201cShares\u201d), subject to the terms of the Company\u2019s "
        "2025 Stock Plan (the \u201cPlan\u201d) and a Common Stock Purchase Agreement to be "
        "executed by you. The purchase price per share will be the fair market value "
        "on the date of grant, as determined by the Board."
    ), space_after=Pt(10))
    add_body(doc, (
        "You should consult with your own tax advisor to determine the tax "
        "consequences of acquiring the Shares."
    ), space_after=Pt(14))

    # Time-Based Vesting
    add_body(doc, "Time-Based Vesting Component", space_after=Pt(6), bold=True)
    add_body(doc, (
        f"{time_shares} of the Shares shall vest on a time basis as follows: "
        f"1/{vest_years} of such shares shall vest on the {cliff}-month anniversary "
        f"of your employment start date, and 1/{vest_months} of such shares shall "
        "vest on each monthly anniversary thereafter, in each case subject to "
        "your continuous service with the Company through the applicable vesting date."
    ), space_after=Pt(14))

    # Milestone-Based Vesting
    add_body(doc, "Milestone-Based Vesting Component", space_after=Pt(6), bold=True)

    milestones = json.loads(args.milestones) if args.milestones else DEFAULT_MILESTONES

    add_body(doc, (
        f"{milestone_shares} of the Shares shall vest upon the completion of the "
        "following milestones. Each milestone vests only once. The Board shall "
        "determine, in its sole discretion, whether a milestone has been achieved. "
        "Vesting is subject to your continuous service through the date of "
        "milestone achievement."
    ), space_after=Pt(10))

    for i, m in enumerate(milestones, 1):
        shares_str = f"{m['shares']:,} shares" if m.get('shares') else "[SHARES TBD]"
        add_body(doc, f"Milestone {i}: {m['description']} \u2014 {shares_str}", space_after=Pt(4))

    # Reserved shares for future milestones
    reserved = getattr(args, 'reserved_shares', None)
    if reserved:
        doc.add_paragraph()
        add_body(doc, (
            f"Reservation for future milestones: {reserved:,} shares. These reserved "
            "shares are not currently assigned to any milestone and will require "
            "separate Board approval and the definition of new milestones before "
            "any vesting can occur. The Company is under no obligation to define "
            "additional milestones or to grant any reserved shares."
        ), space_after=Pt(4))

    doc.add_paragraph()

    # Board Approval
    add_body(doc, "Board Approval", space_after=Pt(6), bold=True)
    add_body(doc, (
        "This Equity Letter reflects management\u2019s recommendation to the Board. "
        "You acknowledge and agree that you have NO guaranteed right to receive "
        "the equity grant described herein until the Board has formally approved "
        "the specific terms of the grant. The Board may, in its sole discretion, "
        "approve, modify, defer, or decline any or all of the equity described herein."
    ), space_after=Pt(14))

    # Forfeiture and Repurchase
    add_body(doc, "Forfeiture and Repurchase", space_after=Pt(6), bold=True)
    add_body(doc, (
        "Upon termination of your Continuous Service (whether voluntary or "
        "involuntary, with or without Cause), all unvested Shares\u2014including "
        "Shares tied to unmet time-based vesting conditions and Shares tied to "
        "milestones not yet achieved\u2014shall be immediately and automatically "
        "forfeited to the Company for no consideration. The Company shall retain "
        "a right of repurchase on all vested Shares, exercisable within 180 days "
        "of the termination of your Continuous Service, at a price equal to the "
        "fair market value as determined by the Board (or the original purchase "
        "price if termination is for Cause, whichever is lower)."
    ), space_after=Pt(14))

    # Clawback
    add_body(doc, "Clawback", space_after=Pt(6), bold=True)
    add_body(doc, (
        "If the Company determines, in its sole discretion, that any Shares were "
        "granted or vested based on fraud, material misrepresentation, breach of "
        "the CIIAA, or any act constituting Cause, the Company may recover such "
        "Shares (or the proceeds from any disposition thereof) from you."
    ), space_after=Pt(14))

    # Milestone Determination
    add_body(doc, "Milestone Determination", space_after=Pt(6), bold=True)
    add_body(doc, (
        "The Board shall determine, in its sole and absolute discretion, whether "
        "a milestone has been achieved. Partial completion of a milestone does "
        "not entitle you to partial vesting of the corresponding Shares. The "
        "Board\u2019s determination shall be final and binding. The Company reserves "
        "the right to modify, replace, or remove milestones at any time upon "
        "written notice to you, provided that milestones previously achieved "
        "and confirmed by the Board shall not be retroactively revoked."
    ), space_after=Pt(14))

    # Legal terms
    add_body(doc, "General Terms", space_after=Pt(6), bold=True)
    add_body(doc, (
        "This Equity Letter constitutes the entire agreement between you and the "
        "Company regarding equity compensation and supersedes any prior equity "
        "provisions in your employment agreements. This Equity Letter does not "
        "alter your at-will employment status. This Equity Letter may not be "
        "amended except by a written agreement signed by both parties. Nothing in "
        "this Equity Letter shall be construed as a guarantee of continued "
        "employment or a limitation on the Company\u2019s right to terminate your "
        "employment at any time, with or without Cause."
    ), space_after=Pt(14))

    # Signature
    add_signature_block(doc, include_employee_name=name)
    return doc


def generate_vesting_schedule_csv(args, output_dir):
    """Generate the milestone vesting schedule as CSV (A1/B1)."""
    name = args.name
    total = args.total_shares or "TBD"
    time_shares = args.time_shares or "TBD"
    milestone_shares = args.milestone_shares or "TBD"
    milestones = json.loads(args.milestones) if args.milestones else DEFAULT_MILESTONES

    safe_name = name.replace(" ", "_")
    csv_path = os.path.join(output_dir, f"LoadGuard_Vesting_Schedule_{safe_name}.csv")

    with open(csv_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow([
            "Tranche", "Item #", "Description", "Shares",
            "% of Total Grant", "Vesting Trigger", "Status"
        ])

        # Time-based rows
        writer.writerow([
            "Time-Based", "T1",
            f"Cliff vesting ({args.cliff_months or 12}-month anniversary)",
            time_shares if isinstance(time_shares, str) else f"{int(time_shares)//4:,}",
            "25% of time tranche",
            f"{args.cliff_months or 12}-month anniversary of start date",
            "Pending"
        ])
        writer.writerow([
            "Time-Based", "T2",
            f"Monthly vesting (months {(args.cliff_months or 12)+1}-{(args.vest_years or 4)*12})",
            "Monthly increments",
            "Remaining 75% of time tranche",
            "Monthly anniversary after cliff",
            "Pending"
        ])

        # Milestone rows
        for i, m in enumerate(milestones, 1):
            shares_str = f"{m['shares']:,}" if m.get('shares') else "TBD"
            pct = ""
            if m.get('shares') and total != "TBD":
                try:
                    pct = f"{m['shares']/int(total)*100:.2f}%"
                except (ValueError, ZeroDivisionError):
                    pct = "TBD"
            writer.writerow([
                "Milestone-Based", f"M{i}", m['description'],
                shares_str, pct or "TBD",
                "Milestone completion (Board determination)",
                "Pending"
            ])

        # Reserved shares for future milestones
        reserved = getattr(args, 'reserved_shares', None)
        if reserved:
            try:
                pct = f"{reserved/int(total)*100:.2f}%" if total != "TBD" else "TBD"
            except (ValueError, ZeroDivisionError):
                pct = "TBD"
            writer.writerow([
                "Reserved", "R1",
                "Reservation for future milestones — requires separate Board approval and new milestone definitions",
                f"{reserved:,}", pct,
                "Future Board action required",
                "Reserved"
            ])

        # Discretionary (only if specified)
        if getattr(args, 'discretionary_bonus', None):
            writer.writerow([
                "Discretionary", "D1",
                "Contributions exceeding expectations that materially improve project success",
                args.discretionary_bonus, args.discretionary_bonus,
                "Board discretion",
                "Pending"
            ])

    return csv_path


def generate_task_schedule(args, is_bonus=False):
    """Generate the task schedule document (A2/B2)."""
    doc = create_doc_with_letterhead()
    name = args.name
    tasks = json.loads(args.tasks) if args.tasks else []

    title_label = "Task Bonus Schedule" if is_bonus else "Task Schedule"

    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"EXHIBIT A")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_label)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'

    add_body(doc, f"Employee: {name}", space_after=Pt(4))
    add_body(doc, f"Date: {today_str()}", space_after=Pt(14))

    # Task table
    comp_label = "Bonus Amount" if is_bonus else "Compensation"
    cols = ["Task #", "Scope / Description", "Deliverable(s)", comp_label,
            "Deadline", "Acceptance Criteria", "Status"]
    num_cols = len(cols)

    if tasks:
        table = doc.add_table(rows=1 + len(tasks), cols=num_cols)
        table.style = 'Table Grid'

        # Header row
        for j, col_name in enumerate(cols):
            cell = table.cell(0, j)
            cell.text = col_name
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'

        # Data rows
        total_comp = 0
        for i, task in enumerate(tasks):
            comp_key = 'bonus' if is_bonus else 'compensation'
            comp_val = task.get(comp_key, task.get('compensation', 'TBD'))
            row_data = [
                str(i + 1),
                task.get('description', 'TBD'),
                task.get('deliverable', 'TBD'),
                comp_val,
                task.get('deadline', 'TBD'),
                task.get('criteria', 'TBD'),
                'Pending'
            ]
            for j, val in enumerate(row_data):
                cell = table.cell(i + 1, j)
                cell.text = val
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Calibri'

            # Sum compensation
            try:
                amt = comp_val.replace('$', '').replace(',', '')
                total_comp += float(amt)
            except (ValueError, AttributeError):
                pass
    else:
        add_body(doc, "[TASKS TO BE DEFINED]", space_after=Pt(14))
        total_comp = 0

    # Terms below table
    doc.add_paragraph()
    total_str = f"${total_comp:,.2f}" if total_comp > 0 else "[TBD]"

    if is_bonus:
        add_body(doc, f"Total Potential Bonus Compensation: {total_str}", space_after=Pt(10), bold=True)
        add_body(doc, (
            "These bonuses are in addition to, and do not replace, the base salary "
            "described in the Offer of Employment Letter."
        ), space_after=Pt(10))
        add_body(doc, (
            "Payment Terms: Task completion bonuses will be paid in the next payroll "
            "cycle following the Company\u2019s confirmation of satisfactory completion of "
            "the deliverable."
        ), space_after=Pt(10))
    else:
        add_body(doc, f"Total Potential Compensation: {total_str}", space_after=Pt(10), bold=True)
        add_body(doc, (
            "Payment Terms: Payment is due Net 15 from the Company\u2019s acceptance of "
            "the deliverable."
        ), space_after=Pt(10))

    add_body(doc, (
        "Deliverable Acceptance: Satisfactory completion of each deliverable shall "
        "be determined by the Chief Executive Officer or their designee."
    ), space_after=Pt(10))
    add_body(doc, (
        f"Amendment: This {title_label} may be amended by mutual written agreement "
        "of the parties."
    ), space_after=Pt(10))
    add_body(doc, (
        "Note: Completion of the tasks listed herein does not guarantee additional "
        "tasks, continued employment, or any other obligation on the part of the Company."
    ), space_after=Pt(14))

    # Signatures
    add_signature_block(doc, include_employee_name=name)
    return doc


def generate_h1b_support_letter(args):
    """Generate the H-1B support letter (B1/B2)."""
    doc = create_doc_with_letterhead()
    name = args.name
    position = args.position

    # From
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(
        f"From: {COMPANY_NAME}\n"
        f"{COMPANY_ADDR}\n"
        f"{COMPANY_CITY}\n"
        f"{today_str()}"
    )
    run.font.size = Pt(11)

    # To
    add_body(doc, (
        "To: U.S. Citizenship and Immigration Services"
    ), space_after=Pt(14))

    # Re
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(f"Re: H-1B Petition for {name} \u2014 {position} (Specialty Occupation)")
    run.bold = True
    run.font.size = Pt(11)

    add_body(doc, "To whom it may concern,", space_after=Pt(12))

    # Company Description
    add_body(doc, (
        f"We submit this letter in support of the H-1B petition for {name} for the "
        f"position of {position} at {COMPANY_NAME}. "
        f"{COMPANY_NAME} is a Delaware C-corporation (EIN: {COMPANY_EIN}) "
        "headquartered in Seattle, Washington. The Company develops advanced "
        "maritime container security technology, including sensor fusion systems, "
        "embedded machine learning, and real-time anomaly detection platforms "
        "for supply chain integrity."
    ), space_after=Pt(14))

    # Specialty Occupation
    add_body(doc, "Specialty Occupation Justification", space_after=Pt(6), bold=True)
    add_body(doc, (
        f"The {position} position at {COMPANY_NAME} is a specialty occupation within "
        "the meaning of 8 C.F.R. \u00a7 214.2(h)(4)(ii). The position requires the "
        "theoretical and practical application of a body of highly specialized "
        "knowledge and the attainment of a bachelor\u2019s degree or higher in a specific "
        "specialty as a minimum for entry into the occupation. A bachelor\u2019s degree "
        "(or its foreign equivalent) in Computer Science, Computer Engineering, "
        "Electrical Engineering, or a directly related field is the minimum "
        "educational requirement."
    ), space_after=Pt(14))

    # Employee Qualifications
    add_body(doc, "Employee Qualifications", space_after=Pt(6), bold=True)
    degree = getattr(args, 'degree', None) or "[DEGREE AND FIELD]"
    university = getattr(args, 'university', None) or "[UNIVERSITY]"
    prior = getattr(args, 'prior_employer', None) or "[PRIOR EXPERIENCE]"
    add_body(doc, (
        f"{name} holds a {degree} from {university}. "
        f"Prior to joining {COMPANY_NAME}, {name} served as {prior}. "
        f"{name}\u2019s specialized education and professional experience provide "
        "the requisite knowledge and skills for this specialty occupation position."
    ), space_after=Pt(14))

    # Prevailing Wage
    add_body(doc, "Prevailing Wage Compliance", space_after=Pt(6), bold=True)
    add_body(doc, (
        f"{COMPANY_NAME} will compensate {name} at or above the prevailing wage "
        "for the position and work location, as determined by the applicable "
        "wage source published by the U.S. Department of Labor. The Company has "
        "filed (or will file) a Labor Condition Application with the DOL and "
        "will comply with all LCA obligations."
    ), space_after=Pt(14))

    # Employment Confirmation
    add_body(doc, "Employment Confirmation", space_after=Pt(6), bold=True)
    is_part_time = getattr(args, 'part_time', False)
    hours_spec = getattr(args, 'hours_per_week', None)
    if is_part_time and hours_spec:
        hours_desc = f"a permanent, part-time position requiring {hours_spec} hours per week"
    elif is_part_time:
        hours_desc = "a permanent, part-time position"
    else:
        hours_desc = "a permanent, full-time position requiring a minimum of 40 hours per week"
    add_body(doc, (
        f"This letter confirms that {COMPANY_NAME} is offering {name} {hours_desc}. "
        f"The position is located at {COMPANY_ADDR}, {COMPANY_CITY}."
    ), space_after=Pt(14))

    # Closing
    add_body(doc, (
        "We respectfully request favorable consideration of this petition. Should "
        "you require any additional information, please do not hesitate to contact "
        "the undersigned."
    ), space_after=Pt(14))

    # Signature (CEO only, no employee acceptance)
    doc.add_paragraph()
    sincerely = doc.add_paragraph()
    sincerely.paragraph_format.space_after = Pt(36)
    run = sincerely.add_run("Sincerely,")
    run.font.size = Pt(11)

    for text, bold in [(CEO_NAME, True), (CEO_TITLE, False), (COMPANY_NAME, False)]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)

    return doc


def generate_specialty_memo(args):
    """Generate the specialty occupation justification memo (B1/B2)."""
    doc = create_doc_with_letterhead()
    name = args.name
    position = args.position

    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INTERNAL MEMORANDUM")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Specialty Occupation Justification")
    run.bold = True
    run.font.size = Pt(12)

    add_body(doc, f"Date: {today_str()}", space_after=Pt(4))
    add_body(doc, f"Re: {name} \u2014 {position}", space_after=Pt(4))
    add_body(doc, "Classification: Confidential \u2014 Attorney-Client Privileged", space_after=Pt(14))

    # Section 1: Position Requirements
    add_body(doc, "1. Position Requirements", space_after=Pt(6), bold=True)
    specialty = getattr(args, 'specialty_justification', None) or "[DESCRIBE WHY THIS ROLE REQUIRES SPECIALIZED KNOWLEDGE]"
    add_body(doc, (
        f"The {position} position at {COMPANY_NAME} requires specialized knowledge "
        "that can only be obtained through completion of a bachelor\u2019s degree program "
        "in a specific specialty field. {specialty}"
    ).format(specialty=specialty), space_after=Pt(14))

    # Section 2: Degree Requirement
    add_body(doc, "2. Degree Requirement", space_after=Pt(6), bold=True)
    add_body(doc, (
        "The minimum educational requirement for this position is a bachelor\u2019s "
        "degree (or its foreign equivalent) in Computer Science, Computer "
        "Engineering, Electrical Engineering, or a directly related field. This "
        "requirement is consistent with industry standards for similar positions "
        "as documented in the Bureau of Labor Statistics Occupational Outlook "
        "Handbook."
    ), space_after=Pt(14))

    # Section 3: Industry Standards
    add_body(doc, "3. Industry Standards", space_after=Pt(6), bold=True)
    add_body(doc, (
        "According to the Bureau of Labor Statistics, positions in software "
        "development, cloud engineering, and artificial intelligence typically "
        "require a bachelor\u2019s degree in computer science or a related field. "
        "The specialized nature of the Company\u2019s work in maritime security "
        "technology, sensor fusion, and embedded machine learning further "
        "necessitates advanced domain-specific education."
    ), space_after=Pt(14))

    # Section 4: Candidate Qualifications
    add_body(doc, "4. Candidate Qualifications", space_after=Pt(6), bold=True)
    degree = getattr(args, 'degree', None) or "[DEGREE AND FIELD]"
    university = getattr(args, 'university', None) or "[UNIVERSITY]"
    prior = getattr(args, 'prior_employer', None) or "[PRIOR EXPERIENCE]"
    add_body(doc, (
        f"{name} holds a {degree} from {university}. "
        f"Prior professional experience: {prior}. "
        f"This combination of education and experience directly qualifies "
        f"{name} for the {position} position and exceeds the minimum "
        "requirements for the specialty occupation."
    ), space_after=Pt(14))

    # Section 5: Conclusion
    add_body(doc, "5. Conclusion", space_after=Pt(6), bold=True)
    add_body(doc, (
        f"Based on the foregoing, the {position} position at {COMPANY_NAME} "
        "meets all four criteria for a specialty occupation under "
        "8 C.F.R. \u00a7 214.2(h)(4)(ii): (1) a bachelor\u2019s degree or higher is "
        "the normal minimum requirement; (2) the degree requirement is common "
        "in the industry; (3) the employer normally requires a degree; and "
        "(4) the nature of the duties is so specialized and complex that the "
        "knowledge required is usually associated with a bachelor\u2019s degree or higher."
    ), space_after=Pt(14))

    return doc


def generate_ciiaa(args):
    """Generate CIIAA by filling in the lawyer template with employee details.

    Clones the template DOCX and fills in:
      - Employee Name / Effective Date header fields
      - Company address on signature page (updated to current HQ)
      - Employee name on signature page and Exhibit A
      - Removes internal lawyer instruction paragraphs
    """
    template_path = os.path.join(
        PROJECT_ROOT, 'Admin', '01. Incorp and Paper',
        'Project LoadGuard Incorporated - Form of Confidential Information '
        'and Invention Assignment Agreement (Employee).docx.docx'
    )
    if not os.path.exists(template_path):
        print(f"  [SKIP] CIIAA template not found: {template_path}")
        return None

    doc = Document(template_path)
    name = args.name
    effective_date = (
        args.start_date
        if (args.start_date and args.start_date.upper() != 'TBD')
        else today_str()
    )

    # ── Remove lawyer instruction paragraphs (first two) ──
    removed = 0
    while removed < 2 and len(doc.paragraphs) > 0:
        first_text = doc.paragraphs[0].text.strip()
        if first_text.startswith('[') and first_text.endswith(']'):
            doc.paragraphs[0]._element.getparent().remove(
                doc.paragraphs[0]._element
            )
            removed += 1
        else:
            break

    # ── Helper: replace text in a paragraph's runs (preserves formatting) ──
    def replace_in_para(para, old, new):
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                return True
        # Fallback: check full paragraph text (single run or merged)
        if old in para.text:
            if para.runs:
                para.runs[0].text = para.text.replace(old, new)
                for run in para.runs[1:]:
                    run.text = ''
                return True
        return False

    # ── Fill in fields ──
    employee_print_filled = False  # track first "(Print Name)" only

    for p in doc.paragraphs:
        text = p.text.strip()

        # Header: Employee Name
        if text == 'Employee Name:':
            replace_in_para(p, 'Employee Name:', f'Employee Name: {name}')

        # Header: Effective Date
        elif text == 'Effective Date:':
            replace_in_para(p, 'Effective Date:', f'Effective Date: {effective_date}')

        # Company address on signature page (update from NJ to current HQ)
        elif '262 Woodfield Road' in text:
            replace_in_para(p, '262 Woodfield Road', COMPANY_ADDR)
            replace_in_para(p, 'Township of Washington, NJ 07676', COMPANY_CITY)

        # Employee print name on signature page (first occurrence only)
        elif text == '(Print Name)' and not employee_print_filled:
            replace_in_para(p, '(Print Name)', name)
            employee_print_filled = True

        # Print Name on Exhibit A
        elif text == 'Print Name of Employee:':
            replace_in_para(p, 'Print Name of Employee:',
                            f'Print Name of Employee: {name}')

    return doc


def generate_arbitration_agreement(args):
    """Generate the Mutual Arbitration Agreement (Attachment B for all variants)."""
    doc = create_doc_with_letterhead()
    name = args.name

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Mutual Arbitration Agreement")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'

    # Body paragraphs — exact language from lawyer template
    body_paras = [
        (
            "To the fullest extent permitted by applicable law, you and "
            f"{COMPANY_NAME} (the \u201cCompany\u201d) agree to arbitrate any and all "
            "claims or disputes relating to or arising from your employment, "
            "including claims by the Company, claims against the Company, and "
            "claims against any parent, affiliate, subsidiary, successor or "
            "predecessor of the Company, and each of the Company\u2019s and these "
            "entities\u2019 respective officers, directors, agents and employees. "
            "To the fullest extent permitted by applicable law, this includes, "
            "but is not limited to, claims of discrimination, harassment, "
            "retaliation, wrongful termination and unfair competition, wage "
            "and hour claims, equity claims, tort claims, contract claims, "
            "common law claims, claims for penalties, and claims based upon "
            "any federal, state or local ordinance, statute, regulation or "
            "constitutional provision."
        ),
        (
            "The parties agree that nothing in this arbitration agreement is "
            "intended to prevent either party from seeking and obtaining "
            "temporary or preliminary injunctive relief in court to prevent "
            "irreparable harm to their confidential information or trade "
            "secrets pending the conclusion of any arbitration."
        ),
        (
            "This arbitration agreement does not apply to claims for "
            "workers\u2019 compensation benefits, unemployment insurance benefits, "
            "or state or federal disability insurance, claims that are subject "
            "to the exclusive jurisdiction of the National Labor Relations "
            "Board, or any other claims that have been expressly excluded from "
            "mandatory arbitration by the Federal Arbitration Act or a "
            "governing law not preempted by the Federal Arbitration Act. If "
            "either party brings both arbitrable and non-arbitrable claims in "
            "the same action or related actions, both agree that the "
            "non-arbitrable claims shall be stayed until the conclusion of "
            "the arbitration, to the fullest extent permitted by law. This "
            "arbitration agreement does not restrict or preclude you from "
            "communicating with, filing an administrative charge or claim "
            "with, or providing testimony to any governmental entity about "
            "any actual or potential violation of law or obtaining relief "
            "through a government agency process, where required by law."
        ),
        (
            "The parties agree that claims shall be resolved on an individual "
            "basis only, and not on a class, collective, or representative "
            "basis on behalf of other employees to the fullest extent "
            "permitted by applicable law (\u201cClass Waiver\u201d). All individual "
            "claims covered by this arbitration agreement shall proceed in "
            "arbitration. Any claim that all or part of the Class Waiver is "
            "invalid, unenforceable, unconscionable, void or voidable may be "
            "determined only by a court. In no case may class, collective or "
            "representative claims proceed in arbitration on behalf of other "
            "employees."
        ),
        (
            "The parties agree that the arbitration shall be conducted by a "
            "single neutral arbitrator through JAMS in accordance with JAMS "
            "Employment Arbitration Rules and Procedures (available at "
            "www.jamsadr.com/rules-employment-arbitration or contact HR for "
            "a copy). To initiate an arbitration, you or the Company must "
            "submit a demand for arbitration to JAMS. Except as to the Class "
            "Waiver and as otherwise required by law, the arbitrator shall "
            "determine arbitrability, including disputes about the formation, "
            "scope, applicability, enforceability or validity of the "
            "arbitration agreement. The Company will bear all JAMS arbitration "
            "fees and administrative costs in excess of the amount of "
            "administrative fees and costs that you otherwise would have been "
            "required to pay if the claims were litigated in court. The "
            "arbitrator shall apply the applicable substantive law in deciding "
            "the claims at issue. Claims will be governed by their applicable "
            "statute of limitations and failure to demand arbitration within "
            "the prescribed time period shall bar the claims as provided by "
            "law. Either party shall have the right to file a motion to "
            "dismiss and/or a motion for summary judgment, which the "
            "arbitrator shall have the authority and obligation to decide by "
            "application of the Federal Rules of Civil Procedure governing "
            "such motions. The parties further agree that either party has "
            "the right to serve the equivalent of an offer of judgment under "
            "Federal Rule of Civil Procedure 68, and if the judgment that the "
            "other party finally obtains in arbitration is not more favorable "
            "than the unaccepted offer, then the other party shall pay the "
            "attorney\u2019s fees and costs incurred after the offer was made to "
            "the same extent that if the case were filed in Court. The parties "
            "understand and agree that the decision or award of the arbitrator "
            "shall be final and binding upon the parties, subject to review "
            "on the grounds set forth in the Federal Arbitration Act. No "
            "arbitration award or decision will have any preclusive effect as "
            "to issues or claims in any dispute with anyone who is not a "
            "named party to the arbitration."
        ),
        (
            "The parties understand and agree that the arbitration of claims "
            "subject to this arbitration agreement shall be instead of a "
            "trial before a court or jury. The parties further understand and "
            "agree that, by entering into this arbitration agreement, they "
            "are expressly waiving any and all rights to a trial before a "
            "court or jury regarding any claims that they now have or which "
            "they may have in the future that are subject to arbitration "
            "under this arbitration agreement."
        ),
        (
            "This arbitration agreement is enforceable under and governed by "
            "the Federal Arbitration Act. In the event that any portion of "
            "this arbitration agreement is held to be invalid or "
            "unenforceable, any such provision shall be severed, and the "
            "remainder of this arbitration agreement will be given full force "
            "and effect. You and the Company understand and agree that this "
            "arbitration agreement contains the complete agreement between "
            "you and the Company regarding the subject matter herein and that "
            "it supersedes any and all prior representations and agreements "
            "between the parties, whether written or oral, on this subject "
            "matter. The provisions of this arbitration agreement shall "
            "survive termination of your employment with the Company. The "
            "Company may offer a new or revised arbitration agreement in "
            "writing by providing notice by email and you will have 30 days "
            "to accept any offered revised agreement by electing to remain "
            "employed and/or accepting the benefits of employment after "
            "receipt of any revised agreement."
        ),
        (
            "By signing below, you acknowledge and agree that you have read "
            "this arbitration agreement carefully, are bound by it and are "
            "WAIVING ANY RIGHT TO HAVE A TRIAL BEFORE A COURT OR JURY OF "
            "ANY AND ALL CLAIMS SUBJECT TO ARBITRATION UNDER THIS ARBITRATION "
            "AGREEMENT."
        ),
    ]

    for para_text in body_paras:
        add_body(doc, para_text, space_after=Pt(10))

    # Signature block — dual column style
    doc.add_paragraph()
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.autofit = True

    # Remove table borders
    tbl = sig_table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        borders = tblPr.find(qn('w:tblBorders'))
        if borders is not None:
            tblPr.remove(borders)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Row 0: signature lines
    for j, text in enumerate(["______________________________", "______________________________"]):
        cell = sig_table.cell(0, j)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    # Row 1: labels
    labels_row1 = [name, f"{CEO_NAME}, CEO"]
    for j, text in enumerate(labels_row1):
        cell = sig_table.cell(1, j)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.bold = True

    # Row 2: "On Behalf of the Company" (right column only)
    cell = sig_table.cell(2, 1)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("On Behalf of the Company")
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.italic = True

    # Row 3: Employee Signature + Date
    cell = sig_table.cell(3, 0)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Employee Signature: ______________________________")
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    cell = sig_table.cell(3, 1)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Date: ______________________________")
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    return doc


# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="LoadGuard Onboarding Document Generator")
    parser.add_argument('--variant', required=True, choices=['A1', 'A2', 'B1', 'B2'])
    parser.add_argument('--name', required=True)
    parser.add_argument('--position', required=True)
    parser.add_argument('--start-date', default='TBD')
    parser.add_argument('--reports-to', default='Chief Executive Officer')
    parser.add_argument('--work-location', default=f'{COMPANY_ADDR}, {COMPANY_CITY}')
    parser.add_argument('--salary', default=None)

    # Equity args (A1/B1)
    parser.add_argument('--total-shares', default=None)
    parser.add_argument('--time-shares', default=None)
    parser.add_argument('--milestone-shares', default=None)
    parser.add_argument('--cliff-months', type=int, default=12)
    parser.add_argument('--vest-years', type=int, default=4)
    parser.add_argument('--milestones', default=None, help="JSON array of milestones")
    parser.add_argument('--reserved-shares', type=int, default=None, help="Shares reserved for future milestones (requires Board approval)")
    parser.add_argument('--discretionary-bonus', default=None, help="Discretionary bonus pool (e.g. 'Up to 2%% equity')")

    # Task args (A2/B2)
    parser.add_argument('--tasks', default=None, help="JSON array of tasks")

    # Immigration args (B1/B2)
    parser.add_argument('--citizenship', default=None)
    parser.add_argument('--visa-status', default=None)
    parser.add_argument('--degree', default=None)
    parser.add_argument('--university', default=None)
    parser.add_argument('--prior-employer', default=None)
    parser.add_argument('--specialty-justification', default=None)

    # Optional
    parser.add_argument('--introductory-period', action='store_true')
    parser.add_argument('--intro-days', type=int, default=30)
    parser.add_argument('--part-time', action='store_true', help="Part-time employment")
    parser.add_argument('--hours-per-week', default=None, help="Hours per week (e.g. '10-20' or '20')")
    parser.add_argument('--offer-only', action='store_true', help="Generate offer letter only (skip supplemental docs)")

    args = parser.parse_args()
    variant = args.variant.upper()
    is_immigration = variant.startswith('B')
    is_equity = variant in ('A1', 'B1')
    is_task = variant in ('A2', 'B2')

    output_dir = ensure_output_dir(args.name)
    safe_name = args.name.replace(" ", "_")
    generated = []

    print(f"\n{'='*60}")
    print(f"  LoadGuard Onboarding — Variant {variant}")
    print(f"  Employee: {args.name}")
    print(f"  Position: {args.position}")
    print(f"  Output:   {output_dir}")
    print(f"{'='*60}\n")

    # 1. Offer Letter (all variants)
    offer_doc = generate_offer_letter(args, immigration=is_immigration, task_based=is_task)
    offer_path = os.path.join(output_dir, f"Offer_of_Employment_{safe_name}.docx")
    offer_doc.save(offer_path)
    generated.append(("Offer of Employment Letter", offer_path))
    print(f"  [OK] Offer Letter: {offer_path}")

    # 2a. CIIAA (all variants — Attachment A)
    ciiaa_doc = generate_ciiaa(args)
    if ciiaa_doc:
        ciiaa_path = os.path.join(output_dir, f"CIIAA_{safe_name}.docx")
        ciiaa_doc.save(ciiaa_path)
        generated.append(("CIIAA (Attachment A)", ciiaa_path))
        print(f"  [OK] CIIAA: {ciiaa_path}")

    # 2b. Arbitration Agreement (all variants — Attachment B)
    arb_doc = generate_arbitration_agreement(args)
    arb_path = os.path.join(output_dir, f"Arbitration_Agreement_{safe_name}.docx")
    arb_doc.save(arb_path)
    generated.append(("Mutual Arbitration Agreement", arb_path))
    print(f"  [OK] Arbitration Agreement: {arb_path}")

    # 3a. Equity documents (A1/B1)
    if is_equity:
        eq_doc = generate_equity_incentive_letter(args)
        eq_path = os.path.join(output_dir, f"Equity_Incentive_Letter_{safe_name}.docx")
        eq_doc.save(eq_path)
        generated.append(("Equity Incentive Letter", eq_path))
        print(f"  [OK] Equity Incentive Letter: {eq_path}")

        csv_path = generate_vesting_schedule_csv(args, output_dir)
        generated.append(("Vesting Schedule (CSV)", csv_path))
        print(f"  [OK] Vesting Schedule: {csv_path}")

    # 3b. Task Schedule (A2/B2) — only if tasks are defined
    if is_task and args.tasks:
        is_bonus = is_immigration  # B2 uses "bonus" language
        task_doc = generate_task_schedule(args, is_bonus=is_bonus)
        label = "Task_Bonus_Schedule" if is_bonus else "Task_Schedule"
        task_path = os.path.join(output_dir, f"{label}_{safe_name}.docx")
        task_doc.save(task_path)
        generated.append((label.replace("_", " "), task_path))
        print(f"  [OK] {label}: {task_path}")

    # 4. Immigration documents (B1/B2) — skip with --offer-only
    if is_immigration and not getattr(args, 'offer_only', False):
        h1b_doc = generate_h1b_support_letter(args)
        h1b_path = os.path.join(output_dir, f"H1B_Support_Letter_{safe_name}.docx")
        h1b_doc.save(h1b_path)
        generated.append(("H-1B Support Letter", h1b_path))
        print(f"  [OK] H-1B Support Letter: {h1b_path}")

        memo_doc = generate_specialty_memo(args)
        memo_path = os.path.join(output_dir, f"Specialty_Occupation_Memo_{safe_name}.docx")
        memo_doc.save(memo_path)
        generated.append(("Specialty Occupation Memo", memo_path))
        print(f"  [OK] Specialty Occupation Memo: {memo_path}")

    # Summary
    print(f"\n{'='*60}")
    print(f"  Generated {len(generated)} document(s):")
    for label, path in generated:
        print(f"    - {label}")
    print(f"{'='*60}\n")

    # Warnings
    warnings = []
    if is_immigration:
        warnings.append("ALL immigration documents should be reviewed by immigration counsel before filing.")
    if not args.salary or args.salary.upper() == "TBD":
        warnings.append("Salary is TBD — must be filled in before signing.")
    if is_equity and not args.total_shares:
        warnings.append("Total shares are TBD — Board must approve equity grant.")
    if is_task and not args.tasks and not getattr(args, "offer_only", False):
        warnings.append("Task schedule is empty — tasks must be defined before signing.")

    if warnings:
        print("  WARNINGS:")
        for w in warnings:
            print(f"    ! {w}")
        print()


if __name__ == '__main__':
    main()
